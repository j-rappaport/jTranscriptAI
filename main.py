import uuid
import os
import re
import struct
import threading
import asyncio
import shutil
import errno
import io
import zipfile
import psycopg2
import psycopg2.extras
import assemblyai as aai
import stripe
import ffmpeg
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from fastapi import FastAPI, UploadFile, File, HTTPException, Header, Request, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from dotenv import load_dotenv
from fastapi import Depends
from typing import List
import mimetypes
import httpx
from jose import jwt as jose_jwt, JWTError
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse

load_dotenv()

stripe.api_key = os.getenv("STRIPE_SECRET_KEY")

CREDIT_PACKS = {
    "30m": {"hours": 0.5, "price_cents": 100,  "label": "30 minutes"},
    "5h":  {"hours": 5,   "price_cents": 1000, "label": "5 hours"},
    "15h": {"hours": 15,  "price_cents": 2500, "label": "15 hours"},
    "50h": {"hours": 50,  "price_cents": 7000, "label": "50 hours"},
}

app = FastAPI(title="jTranscript")

@app.on_event("startup")
async def mark_stuck_jobs():
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE jobs SET status='error', error='Interrupted — server restarted during transcription' "
                "WHERE status IN ('pending', 'transcribing')"
            )

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_methods=["*"],
    allow_headers=["*"],
)

_jwks_cache = None

def _get_jwks():
    global _jwks_cache
    if _jwks_cache is None:
        _jwks_cache = httpx.get(os.getenv("CLERK_JWKS_URL")).json()
    return _jwks_cache

def _verify_token(token: str) -> str:
    global _jwks_cache
    jwks = _get_jwks()
    try:
        kid = jose_jwt.get_unverified_header(token).get("kid")
        key = next((k for k in jwks.get("keys", []) if k.get("kid") == kid), None)
        if key is None:
            _jwks_cache = None
            jwks = _get_jwks()
            key = next((k for k in jwks.get("keys", []) if k.get("kid") == kid), None)
        if key is None:
            raise HTTPException(status_code=401, detail="Unknown signing key")
        payload = jose_jwt.decode(token, key, algorithms=["RS256"], options={"verify_aud": False})
        return payload["sub"]
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid token")
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=401, detail="Authentication error")

def require_auth(authorization: str = Header(default=None)) -> str:
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Not authenticated",
                            headers={"WWW-Authenticate": "Bearer"})
    return _verify_token(authorization[7:])

aai.settings.api_key = os.getenv("ASSEMBLYAI_API_KEY")


# ---------------------------------------------------------------------------
# Database helpers
# ---------------------------------------------------------------------------

def get_conn():
    return psycopg2.connect(os.getenv("DATABASE_URL"))


def db_create_job(job_id: str, filename: str, audio_path: str, user_id: str):
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                INSERT INTO jobs (id, status, filename, audio_path, user_id)
                VALUES (%s, %s, %s, %s, %s)
                """,
                (job_id, "pending", filename, audio_path, user_id)
            )


def db_update_status(job_id: str, status: str, error: str = None):
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE jobs SET status=%s, error=%s WHERE id=%s",
                (status, error, job_id)
            )


def db_update_done(job_id: str, blocks: list, duration_ms: int):
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                UPDATE jobs
                SET status='done', utterances=%s, audio_duration_ms=%s
                WHERE id=%s
                """,
                (psycopg2.extras.Json(blocks), duration_ms, job_id)
            )


def db_get_job(job_id: str) -> dict:
    with get_conn() as conn:
        with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute("SELECT * FROM jobs WHERE id=%s", (job_id,))
            row = cur.fetchone()
            return dict(row) if row else None

def db_update_blocks(job_id: str, blocks: list):
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE jobs SET utterances=%s WHERE id=%s",
                (psycopg2.extras.Json(blocks), job_id)
            )


def db_get_all_jobs(user_id: str) -> list:
    with get_conn() as conn:
        with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute(
                "SELECT id, status, filename, audio_duration_ms, created_at, audio_path FROM jobs WHERE user_id=%s ORDER BY created_at DESC",
                (user_id,)
            )
            rows = [dict(row) for row in cur.fetchall()]
            for row in rows:
                row["audio_available"] = bool(row["audio_path"] and os.path.exists(row["audio_path"]))
                del row["audio_path"]
            return rows


def db_get_credits(user_id: str) -> int:
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT credits_ms FROM user_credits WHERE user_id=%s", (user_id,))
            row = cur.fetchone()
            return row[0] if row else 0

def db_add_credits(user_id: str, ms: int):
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                INSERT INTO user_credits (user_id, credits_ms)
                VALUES (%s, %s)
                ON CONFLICT (user_id) DO UPDATE
                SET credits_ms = user_credits.credits_ms + EXCLUDED.credits_ms
                """,
                (user_id, ms)
            )

def db_deduct_credits(user_id: str, ms: int):
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE user_credits SET credits_ms = credits_ms - %s WHERE user_id = %s",
                (ms, user_id)
            )


def read_tutc_timestamp(data: bytes):
    idx = data.find(b'TUTC')
    if idx >= 0 and idx + 16 <= len(data):
        return struct.unpack('<Q', data[idx+8:idx+16])[0]
    return None

def merge_audio_files(paths: list, output_path: str):
    if len(paths) == 1:
        (ffmpeg.input(paths[0]).audio
         .output(output_path, acodec='libmp3lame', audio_bitrate='192k')
         .run(overwrite_output=True, quiet=True))
    else:
        inputs = [ffmpeg.input(p) for p in paths]
        (ffmpeg.concat(*[i.audio for i in inputs], v=0, a=1)
         .output(output_path, acodec='libmp3lame', audio_bitrate='192k')
         .run(overwrite_output=True, quiet=True))

def extract_case_name(filename: str) -> str:
    m = re.match(r'^(.+?)_\d{8}-\d{4}_', filename)
    return m.group(1) if m else os.path.splitext(filename)[0]


def delete_old_audio(user_id: str):
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT audio_path FROM jobs WHERE audio_path IS NOT NULL AND user_id=%s", (user_id,))
            rows = cur.fetchall()
            for row in rows:
                path = row[0]
                if path and os.path.exists(path):
                    os.remove(path)
# ---------------------------------------------------------------------------
# Docx generation
# ---------------------------------------------------------------------------

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "templates", "Transcript.dotm")
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_ABBREV_RE = re.compile(r"\b(Mr|Mrs|Ms|Dr|Jr|Sr|vs|etc|No|St)\.  ", re.IGNORECASE)

def _open_dotm():
    with open(TEMPLATE_PATH, "rb") as f:
        data = f.read()
    inp = zipfile.ZipFile(io.BytesIO(data))
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as out:
        for item in inp.infolist():
            content = inp.read(item.filename)
            if item.filename == "[Content_Types].xml":
                content = content.replace(
                    b"application/vnd.ms-word.template.macroEnabledTemplate.main+xml",
                    b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
                )
            out.writestr(item, content)
    buf.seek(0)
    return Document(buf)

def _normalize(text):
    text = re.sub(r"([.?])\s*", r"\1  ", text)
    text = _ABBREV_RE.sub(r"\1. ", text)
    return text.rstrip()

def _compute_roles(blocks):
    qa_on, next_role, roles = False, "Q", []
    for b in blocks:
        if b.get("type") == "qa_toggle":
            qa_on = not qa_on
            if qa_on:
                next_role = "Q"
            roles.append("")
        elif b.get("type") == "utterance" and qa_on:
            r = next_role
            next_role = "A" if r == "Q" else "Q"
            roles.append(r)
        else:
            roles.append("")
    return roles

def _add_para(doc, text=""):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.9
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if text:
        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(12)
    return p

def generate_docx(blocks):
    doc = _open_dotm()

    # Keep first 4 preamble paragraphs, drop the rest
    body = doc.element.body
    paras = body.findall(f"{{{_W_NS}}}p")
    for p in paras[4:]:
        body.remove(p)

    _add_para(doc)  # blank separator

    roles = _compute_roles(blocks)
    prev_role = ""

    for i, block in enumerate(blocks):
        btype = block.get("type", "utterance")
        role = roles[i]
        fields = block.get("fields", {})

        if btype == "qa_toggle":
            continue

        if btype == "utterance":
            text = _normalize(block.get("text", ""))
            if role == "Q" and prev_role == "":
                _add_para(doc, f"BY {block['speaker']}:")
            line = f"\t{role}\t{text}" if role else f"\t\t{block['speaker']}:  {text}"
            _add_para(doc, line)
            prev_role = role
            continue

        # Meta blocks always break the Q/A "BY" header chain
        prev_role = ""

        if btype in ("witness_sworn", "witness_sworn_gj"):
            name = (fields.get("NAME") or "").strip().upper() or "______"
            boilerplate = (
                "Was thereupon called as a witness on behalf of the State; and, having been first duly sworn, was examined and testified as follows:"
                if btype == "witness_sworn"
                else "Was thereupon called as a witness; and, having been first duly sworn, was examined and testified as follows:"
            )
            p = _add_para(doc)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(name)
            run.font.name = "Courier New"
            run.font.size = Pt(12)
            run.underline = True
            _add_para(doc, boilerplate)

        elif btype == "pause_in_proceedings":
            start, end = fields.get("START", "___"), fields.get("END", "___")
            p = _add_para(doc, f"(Pause in proceedings, {start} - {end})")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif btype == "whispered_discussion":
            start, end = fields.get("START", "___"), fields.get("END", "___")
            p = _add_para(doc, f"(Whispered discussion, off the record, {start} - {end})")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif btype == "exhibit_received":
            role_f = fields.get("ROLE", "___")
            number = fields.get("NUMBER", "___")
            p = _add_para(doc, f"({role_f}'s Exhibit No. {number} received.)")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif btype == "header":
            text = fields.get("TEXT", "")
            if text:
                p = _add_para(doc)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.font.name = "Courier New"
                run.font.size = Pt(12)
                run.underline = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------


@app.post("/jobs")
async def create_job(files: List[UploadFile] = File(...), word_boost: str = Form(default=""), user: str = Depends(require_auth)):
    if db_get_credits(user) <= 0:
        raise HTTPException(status_code=402, detail="No credits remaining. Please purchase more to continue.")

    audio_files = [f for f in files if not f.filename.lower().endswith('.trs')]
    if not audio_files:
        raise HTTPException(status_code=400, detail="No audio files provided.")

    free_bytes = shutil.disk_usage(UPLOAD_DIR).free
    if free_bytes < 150 * 1024 * 1024:
        raise HTTPException(status_code=507, detail="Server storage is full — contact administrator.")

    delete_old_audio(user)
    job_id = str(uuid.uuid4())

    needs_ffmpeg = len(audio_files) > 1 or any(f.filename.lower().endswith('.trm') for f in audio_files)

    try:
        if needs_ffmpeg:
            file_data = []
            for i, f in enumerate(audio_files):
                content = await f.read()
                ext = os.path.splitext(f.filename)[1]
                temp_path = os.path.join(UPLOAD_DIR, f"{job_id}_{i}{ext}")
                with open(temp_path, "wb") as fp:
                    fp.write(content)
                ts = read_tutc_timestamp(content[:512])
                file_data.append((ts if ts is not None else float('inf'), f.filename, temp_path))

            file_data.sort(key=lambda x: (x[0], x[1]))
            sorted_paths = [d[2] for d in file_data]

            audio_path = os.path.join(UPLOAD_DIR, f"{job_id}.mp3")
            await asyncio.to_thread(merge_audio_files, sorted_paths, audio_path)

            for p in sorted_paths:
                if os.path.exists(p):
                    os.remove(p)

            filename = extract_case_name(file_data[0][1])
        else:
            f = audio_files[0]
            content = await f.read()
            ext = os.path.splitext(f.filename)[1]
            audio_path = os.path.join(UPLOAD_DIR, f"{job_id}{ext}")
            with open(audio_path, "wb") as fp:
                fp.write(content)
            filename = f.filename
    except OSError as e:
        if e.errno == errno.ENOSPC:
            raise HTTPException(status_code=507, detail="Server storage is full — contact administrator.")
        raise

    db_create_job(job_id, filename, audio_path, user)

    boost_words = [w.strip() for w in word_boost.splitlines() if w.strip()]

    threading.Thread(
        target=run_transcription,
        args=(job_id, audio_path, user, boost_words),
        daemon=True
    ).start()

    return {"job_id": job_id, "status": "pending"}


@app.get("/jobs/{job_id}")
def get_job(job_id: str, user: str = Depends(require_auth)):
    job = db_get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    job["blocks"] = job.pop("utterances")
    return job


@app.get("/jobs/{job_id}/audio")
def get_audio(job_id: str):
    job = db_get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    path = job["audio_path"]
    mime_type, _ = mimetypes.guess_type(path)
    print(f"Serving audio: {path}, mime: {mime_type}")
    return FileResponse(path, media_type=mime_type or "audio/mpeg")

@app.get("/jobs")
def get_all_jobs(user: str = Depends(require_auth)):
    return db_get_all_jobs(user)


@app.put("/jobs/{job_id}/blocks")
def update_blocks(job_id: str, body: dict, user: str = Depends(require_auth)):
    job = db_get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    db_update_blocks(job_id, body["blocks"])
    return {"ok": True}


@app.get("/jobs/{job_id}/docx")
def get_docx(job_id: str, user: str = Depends(require_auth)):
    job = db_get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    blocks = job.get("utterances") or []
    base = os.path.splitext(job.get("filename", "transcript"))[0]
    buf = generate_docx(blocks)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{base}_transcript.docx"'},
    )


@app.get("/jobs/{job_id}/audio-available")
def audio_available(job_id: str, user: str = Depends(require_auth)):
    job = db_get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    path = job["audio_path"]
    return {"available": bool(path and os.path.exists(path))}


@app.get("/credits")
def get_credits(user: str = Depends(require_auth)):
    ms = db_get_credits(user)
    return {"credits_ms": ms, "credits_hours": round(ms / 3_600_000, 2)}


@app.post("/billing/checkout")
async def create_checkout(body: dict, user: str = Depends(require_auth)):
    pack = CREDIT_PACKS.get(body.get("pack"))
    if not pack:
        raise HTTPException(status_code=400, detail="Invalid pack")
    app_url = os.getenv("APP_URL", "http://localhost:5173")
    session = stripe.checkout.Session.create(
        payment_method_types=["card"],
        line_items=[{
            "price_data": {
                "currency": "usd",
                "product_data": {"name": f"jTranscript — {pack['label']} transcription credit"},
                "unit_amount": pack["price_cents"],
            },
            "quantity": 1,
        }],
        mode="payment",
        success_url=f"{app_url}/?payment=success",
        cancel_url=f"{app_url}/?payment=cancelled",
        client_reference_id=user,
        metadata={"user_id": user, "credits_ms": str(int(pack["hours"] * 3_600_000))},
    )
    return {"url": session.url}


@app.post("/billing/webhook")
async def stripe_webhook(request: Request):
    payload = await request.body()
    sig = request.headers.get("stripe-signature")
    try:
        event = stripe.Webhook.construct_event(
            payload, sig, os.getenv("STRIPE_WEBHOOK_SECRET")
        )
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid webhook signature")
    if event["type"] == "checkout.session.completed":
        session = event["data"]["object"]
        user_id = session["metadata"]["user_id"]
        credits_ms = int(session["metadata"]["credits_ms"])
        db_add_credits(user_id, credits_ms)
    return {"ok": True}

# ---------------------------------------------------------------------------
# Transcription
# ---------------------------------------------------------------------------

def run_transcription(job_id: str, audio_path: str, user_id: str, boost_words: list = []):
    try:
        db_update_status(job_id, "transcribing")

        config = aai.TranscriptionConfig(
            speaker_labels=True,
            speech_models=["universal-3-pro", "universal-2"],
            word_boost=boost_words if boost_words else None,
            # prompt=(
            #     "Always: Transcribe speech exactly as heard. If uncertain or audio is unclear, mark as (indiscernible). "
            #     "After the first output, review the transcript again. Pay close attention to hallucinations, misspellings, or errors, "
            #     "and revise them like a computer performing spell and grammar checks. "
            #     "Ensure words and phrases make grammatical sense in sentences. "
            #     "When a voice says \"Okay\" as a brief acknowledgment or interruption of another speaker, treat it as the start of a new speaker turn."
            # ),
        )
        transcriber = aai.Transcriber()
        transcript = transcriber.transcribe(audio_path, config=config)

        if transcript.status == aai.TranscriptStatus.error:
            raise RuntimeError(transcript.error)

        blocks = [
            {
                "type": "utterance",
                "speaker": f"SPEAKER {u.speaker}",
                "text": u.text,
                "start_ms": u.start,
                "end_ms": u.end
            }
            for u in transcript.utterances
        ]

        duration_ms = int(transcript.audio_duration * 1000)
        db_update_done(job_id, blocks, duration_ms)
        db_deduct_credits(user_id, duration_ms)

    except Exception as e:
        db_update_status(job_id, "error", str(e))

FRONTEND_DIST = os.path.join(os.path.dirname(__file__), "frontend", "dist")
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

app.mount("/assets", StaticFiles(directory=os.path.join(FRONTEND_DIST, "assets")), name="assets")

@app.get("/")
async def serve_root():
    with open(os.path.join(FRONTEND_DIST, "index.html")) as f:
        return HTMLResponse(f.read())

@app.get("/{full_path:path}", response_class=HTMLResponse)
async def serve_frontend(full_path: str):
    with open(os.path.join(FRONTEND_DIST, "index.html")) as f:
        return HTMLResponse(f.read())